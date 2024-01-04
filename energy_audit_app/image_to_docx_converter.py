from typing import List, Tuple

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm


class DOCXConverter:
    def __init__(self) -> None:
        self.doc = Document()
        self.configure_default_style()

    def convert(
        self,
        conversion_file_path: str,
        graph_data: List[Tuple[str, List[Tuple[str, str, str]]]],
    ) -> None:
        self.change_orientation_and_margins()

        section_number = 1
        for section_title, graphs in graph_data:
            if section_title:
                self.add_section_title(section_title)

            graph_number = 1
            for image_file, graph_name, comment in graphs:
                figure_title = f"Рис. {section_number}.{graph_number}. {graph_name}"
                self.add_picture_with_text(image_file, figure_title, comment)
                graph_number += 1

            section_number += 1
            self.doc.add_section(WD_SECTION.NEW_PAGE)

        self.doc.save(conversion_file_path)

    def add_section_title(self, title: str) -> None:
        if title:
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run(title)
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.paragraph_format.space_after = Pt(12)

    def add_picture_with_text(
        self, image_file: str, graph_name: str, comment: str
    ) -> None:
        table = self.doc.add_table(rows=2, cols=1)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.style = "Table Grid"

        cell_img = table.cell(0, 0)
        paragraph_img = cell_img.paragraphs[0]
        run_img = paragraph_img.add_run()
        run_img.add_picture(image_file, width=self.doc.sections[0].page_width / 1.4)

        cell_title = table.cell(1, 0)
        paragraph_title = cell_title.paragraphs[0]
        paragraph_title.add_run(graph_name)
        paragraph_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        comment_paragraph = self.doc.add_paragraph(comment)
        comment_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def change_orientation_and_margins(self) -> None:
        for section in self.doc.sections:
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height

            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            section.left_margin = Cm(2.0)

    def configure_default_style(self) -> None:
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(14)

        paragraph_format = style.paragraph_format
        paragraph_format.first_line_indent = None
        paragraph_format.space_after = Pt(8)
        paragraph_format.line_spacing = 1

    @staticmethod
    def _get_image(image_file: str) -> Tuple[str, int, int]:
        img = Image.open(image_file)
        width, height = img.size
        return image_file, width, height
